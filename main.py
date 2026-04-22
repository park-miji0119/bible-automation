"""매주 성경 시나리오를 생성하고 네이버 메일로 발송합니다.

환경변수:
    ANTHROPIC_API_KEY  - Claude API 키
    NAVER_EMAIL        - 발송/수신 네이버 계정 (예: xxx@naver.com)
    NAVER_APP_PASSWORD - 네이버 메일 SMTP 앱 비밀번호
    MAIL_RECIPIENT     - 수신자 (미지정 시 NAVER_EMAIL 로 발송)
    WEEK_OFFSET        - (선택) 주제 순환 오프셋 정수
    TOPIC_IDS          - (선택) 1-indexed 번호 목록. 쉼표/공백 구분. 예: "8,44"
                          지정 시 해당 번호들의 시나리오를 모두 생성해서
                          하나의 워드 파일 / 한 통의 메일로 묶어 발송.
                          미지정 시 현재 주차의 연속된 2개 주제를 자동 선택.
"""

from __future__ import annotations

import os
import re
import smtplib
import sys
from datetime import datetime, timezone, timedelta
from email.message import EmailMessage
from pathlib import Path

from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor

from prompt_template import SYSTEM_PROMPT, build_prompt
from topics import TOPICS

MODEL = "claude-opus-4-5"
KST = timezone(timedelta(hours=9))

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass


def sanitize_text(text: str) -> str:
    # Claude 출력에 간혹 섞이는 U+2028 / U+2029 / U+FEFF 만 정상 개행·제거로 치환.
    # 일반 공백(0x20)은 절대 건드리지 않는다 — 과거 invisible char 이 공백으로 바뀌어
    # 단어 사이마다 줄바꿈이 들어가는 사고가 있었음. 앞으로는 \u escape 로 고정.
    return (
        text.replace("\u2028", "\n")
            .replace("\u2029", "\n\n")
            .replace("\ufeff", "")
    )


def current_week_index() -> int:
    now = datetime.now(KST)
    iso_year, iso_week, _ = now.isocalendar()
    base = iso_year * 53 + iso_week
    offset = int(os.environ.get("WEEK_OFFSET", "0"))
    return base + offset


def generate_scenario(topic: str, narrator: str, reference: str) -> str:
    client = Anthropic()
    message = client.messages.create(
        model=MODEL,
        max_tokens=12000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": build_prompt(topic, narrator, reference)}],
    )
    raw = "".join(block.text for block in message.content if block.type == "text")
    return sanitize_text(raw)


def safe_filename(value: str) -> str:
    return re.sub(r"[^\w\-._ ]", "_", value).strip()[:60] or "scenario"


_SCENE_PATTERN = re.compile(r"^\[(?:장면|Scene)\s*:", re.IGNORECASE)
_DIALOG_PATTERN = re.compile(r'^["“].+["”]\s*$')


def _render_paragraph(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph()
        return

    if _SCENE_PATTERN.match(stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)
        return

    if _DIALOG_PATTERN.match(stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.bold = True
        return

    doc.add_paragraph(line)


def _render_scenario(doc: Document, scenario: str, level_shift: int = 0) -> None:
    """scenario 본문(마크다운 형태)을 docx 에 렌더링. level_shift 는 # 단계 내림폭."""
    for line in scenario.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=min(1 + level_shift, 9))
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=min(2 + level_shift, 9))
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=min(3 + level_shift, 9))
        elif stripped == "---":
            doc.add_paragraph().add_run("─" * 40)
        else:
            _render_paragraph(doc, line)


def save_docx(items: list[tuple[str, str, str, str]], out_dir: Path) -> Path:
    """items: [(topic, narrator, reference, scenario), ...] 를 한 워드 파일로 저장."""
    if not items:
        raise ValueError("save_docx requires at least one item")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    today_kst = datetime.now(KST).strftime("%Y-%m-%d")

    if len(items) == 1:
        topic, narrator, reference, scenario = items[0]
        heading = doc.add_heading(topic, level=0)
        heading.alignment = 1
        meta = doc.add_paragraph()
        meta.add_run(f"화자: {narrator}    본문: {reference}").italic = True
        doc.add_paragraph(f"생성일: {today_kst} (KST)")
        doc.add_paragraph()
        _render_scenario(doc, scenario)
    else:
        title = doc.add_heading(f"{today_kst} 주간 성경 시나리오 · {len(items)}편", level=0)
        title.alignment = 1
        toc = doc.add_paragraph()
        toc.add_run("수록 주제: ").bold = True
        toc.add_run(" · ".join(f"{i+1}) {topic}" for i, (topic, *_rest) in enumerate(items)))
        doc.add_paragraph(f"생성일: {today_kst} (KST)")

        for i, (topic, narrator, reference, scenario) in enumerate(items):
            if i > 0:
                doc.add_page_break()
            doc.add_heading(f"{i+1}. {topic}", level=1)
            meta = doc.add_paragraph()
            meta.add_run(f"화자: {narrator}    본문: {reference}").italic = True
            doc.add_paragraph()
            _render_scenario(doc, scenario, level_shift=1)

    out_dir.mkdir(parents=True, exist_ok=True)
    first_topic = items[0][0]
    if len(items) == 1:
        filename = f"{today_kst}_{safe_filename(first_topic)}.docx"
    else:
        filename = f"{today_kst}_{safe_filename(first_topic)}_외{len(items)-1}편.docx"
    path = out_dir / filename
    doc.save(path)
    return path


def _attachment_name(docx_path: Path) -> str:
    today_kst = datetime.now(KST).strftime("%Y-%m-%d")
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", docx_path.stem).strip("_")
    if slug and slug != today_kst:
        base = slug
    else:
        base = f"{today_kst}_bible_scenario"
    if not base.lower().endswith(".docx"):
        base += ".docx"
    return base


def send_email(docx_path: Path, items: list[tuple[str, str, str, str]]) -> None:
    """한 워드 파일에 items 전체가 담겨있다고 가정하고 한 통의 메일로 발송."""
    sender = os.environ["NAVER_EMAIL"]
    password = os.environ["NAVER_APP_PASSWORD"]
    recipient = os.environ.get("MAIL_RECIPIENT", sender)

    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = recipient
    today_kst = datetime.now(KST).strftime("%Y-%m-%d")

    if len(items) == 1:
        topic, narrator, _, _ = items[0]
        msg["Subject"] = f"[주간 성경 시나리오] {today_kst} · {topic} ({narrator})"
    else:
        narrators = " · ".join(narrator for _, narrator, _, _ in items)
        msg["Subject"] = f"[주간 성경 시나리오] {today_kst} · {len(items)}편 ({narrators})"

    lines = [
        "안녕하세요.",
        "",
        f"이번 주 성경 스토리텔링 시나리오 {len(items)}편을 전달드립니다.",
        "",
    ]
    for i, (topic, narrator, reference, _scenario) in enumerate(items, start=1):
        lines.append(f"  {i}. {topic}")
        lines.append(f"     화자: {narrator}  |  본문: {reference}")
    lines += [
        "",
        "전체 내용은 첨부된 워드 파일을 확인해 주세요.",
        "— Bible Automation",
    ]
    msg.set_content("\n".join(lines))

    attach_name = _attachment_name(docx_path)
    msg.add_attachment(
        docx_path.read_bytes(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=attach_name,
    )

    with smtplib.SMTP_SSL("smtp.naver.com", 465, timeout=30) as smtp:
        smtp.login(sender, password)
        smtp.send_message(msg)


def parse_topic_ids(raw: str) -> list[int]:
    """쉼표/공백 구분 1-indexed 번호를 0-indexed 리스트로 변환. 범위 밖은 무시."""
    ids: list[int] = []
    for tok in re.split(r"[\s,]+", raw.strip()):
        if not tok:
            continue
        try:
            n = int(tok)
        except ValueError:
            print(f"  skip non-integer token: {tok!r}")
            continue
        if 1 <= n <= len(TOPICS):
            ids.append(n - 1)
        else:
            print(f"  skip out-of-range id: {n} (valid 1..{len(TOPICS)})")
    return ids


def run_batch(picks: list[tuple[int, tuple[str, str, str]]], out_dir: Path) -> None:
    """picks: [(1-indexed-id, (topic, narrator, reference)), ...]
    각 picks 에 대해 시나리오를 생성 후 한 docx 로 묶어 한 통의 메일로 발송."""
    items: list[tuple[str, str, str, str]] = []
    for idx1, (topic, narrator, reference) in picks:
        print(f"[#{idx1}] topic={topic} / narrator={narrator} / ref={reference}")
        scenario = generate_scenario(topic, narrator, reference)
        print(f"       scenario generated: {len(scenario)} chars")
        items.append((topic, narrator, reference, scenario))

    docx_path = save_docx(items, out_dir)
    print(f"saved: {docx_path}")
    send_email(docx_path, items)
    print(f"email sent: {len(items)} scenario(s) in one docx.")


def weekly_picks(week_index: int, count: int = 2) -> list[tuple[int, tuple[str, str, str]]]:
    """주차 인덱스 기준으로 연속된 count 개 주제를 비중복으로 선택.
    batch_k = week_index 이면 topic slice = TOPICS[count*k .. count*k+count-1]."""
    start = (count * week_index) % len(TOPICS)
    result: list[tuple[int, tuple[str, str, str]]] = []
    for i in range(count):
        idx0 = (start + i) % len(TOPICS)
        result.append((idx0 + 1, TOPICS[idx0]))
    return result


def main() -> int:
    out_dir = Path(__file__).parent / "output"

    raw_ids = os.environ.get("TOPIC_IDS", "").strip()
    if raw_ids:
        ids = parse_topic_ids(raw_ids)
        if not ids:
            print(f"TOPIC_IDS set ({raw_ids!r}) but yielded no valid ids; aborting.")
            return 1
        picks = [(idx0 + 1, TOPICS[idx0]) for idx0 in ids]
        print(f"[manual] {len(picks)} topic(s) selected via TOPIC_IDS={raw_ids!r}")
    else:
        week_index = current_week_index()
        picks = weekly_picks(week_index, count=2)
        print(f"[weekly] week={week_index}, auto-picked {len(picks)} consecutive topics")

    run_batch(picks, out_dir)
    return 0


if __name__ == "__main__":
    sys.exit(main())
